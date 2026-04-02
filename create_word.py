#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera un file Word (.docx) con il contenuto del sito Evolution Engineering
Usa python-docx: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colori
NAVY = RGBColor(10, 25, 47)
BLUE = RGBColor(0, 129, 201)
GRAY = RGBColor(107, 114, 128)

def set_cell_background(cell, fill_color):
    """Imposta il colore di sfondo di una cella"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_text_color(paragraph, color):
    """Imposta il colore del testo in una cella"""
    for run in paragraph.runs:
        run.font.color.rgb = color

# Crea il documento
doc = Document()

# Imposta i margini
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ===== COPERTINA =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EVOLUTION ENGINEERING & SERVICES S.R.L.")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Soluzioni Metalliche Integrate")
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Carpenteria Strutturale | Piping Industriale | Moduli Prefabbricati")
run.font.size = Pt(11)
run.font.color.rgb = GRAY

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(48)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Documento Word per Revisione\nModifica, salva e rispedisci per implementare le correzioni")
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = GRAY

doc.add_page_break()

# ===== CHI SIAMO =====
h1 = doc.add_heading("CHI SIAMO", level=1)
h1_format = h1.runs[0]
h1_format.font.color.rgb = NAVY

p = doc.add_paragraph()
run = p.add_run("Evolution Engineering & Services S.r.l. ")
run.font.bold = True
run = p.add_run("è un general contractor specializzato in soluzioni metalliche integrate per EPC contractor e system integrator europei.")

p = doc.add_paragraph()
run = p.add_run("40+ anni di esperienza ")
run.font.bold = True
run = p.add_run("nella realizzazione di carpenteria strutturale, piping industriale e moduli prefabbricati per i settori energia, oil & gas, farmaceutico e navale.")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
shading_elm = OxmlElement('w:shd')
shading_elm.set(qn('w:fill'), 'F0F4F8')
p._element.get_or_add_pPr().append(shading_elm)
run = p.add_run("Certificazioni in corso: ISO 9001:2015 e EN 9100:2018")
run.font.bold = True

# ===== SERVIZI =====
h2 = doc.add_heading("SERVIZI PRINCIPALI", level=2)
h2_format = h2.runs[0]
h2_format.font.color.rgb = BLUE

p = doc.add_paragraph()
run = p.add_run("Carpenteria Metallica")
run.font.bold = True
p.add_run(": Strutture in acciaio S275/S355, travi, profili, bullonature. Lavorazioni conformi a EN 1090, EN ISO 3834.")

p = doc.add_paragraph()
run = p.add_run("Piping Industriale")
run.font.bold = True
p.add_run(": Prefabbricazione e montaggio tubi acciaio/inox. DN25-DN500, saldatura qualificata EN ISO 15614, controlli NDT completi.")

p = doc.add_paragraph()
run = p.add_run("Skid & Moduli Prefabbricati")
run.font.bold = True
p.add_run(": Moduli process prefabbricati, strutture modulari per settori GMP e oil & gas.")

p = doc.add_paragraph()
run = p.add_run("Progettazione")
run.font.bold = True
p.add_run(": CAD 3D, dimensionamento secondo normative internazionali (ASME, EN, PED).")

p = doc.add_paragraph()
run = p.add_run("Controlli NDT")
run.font.bold = True
p.add_run(": Radiografie (RT), ultrasuoni (UT), liquidi penetranti (PT), particelle magnetiche (MT). Personale certificato EN ISO 9712.")

# ===== SETTORI =====
h2 = doc.add_heading("SETTORI DI APPLICAZIONE", level=2)
h2_format = h2.runs[0]
h2_format.font.color.rgb = BLUE

sectors = [
    ("Energia e Cogenerazione", "Impianti termici, boiler, economizzatori"),
    ("Oil & Gas", "Topsides, manifold, skid process"),
    ("Farmaceutico (GMP)", "Impianti sterili, classificati, tracciabili"),
    ("Navale", "Scafi, sezioni di coperta, componenti strutturali"),
    ("Industria Pesante", "Infrastrutture, grandi strutture metalliche"),
]

for sector_name, sector_desc in sectors:
    p = doc.add_paragraph()
    run = p.add_run(f"{sector_name}: ")
    run.font.bold = True
    p.add_run(sector_desc)

doc.add_page_break()

# ===== CAPACITÀ PRODUTTIVE =====
h1 = doc.add_heading("CAPACITÀ PRODUTTIVE", level=1)
h1_format = h1.runs[0]
h1_format.font.color.rgb = NAVY

# Crea tabella
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Attrezzatura"
hdr_cells[1].text = "Specifiche"

# Colore header
for cell in hdr_cells:
    set_cell_background(cell, "0A192F")
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

# Dati attrezzature
equipment_data = [
    ("Taglio Plasma CNC", "Spessori 1-80mm | Piano 3000×12000mm | Tolleranza ±0,5mm"),
    ("Taglio Ossigas CNC", "Spessori 10-250mm | 3 cannelli | Piano 4000×12000mm"),
    ("Saldatura SAW", "Corrente fino a 1200A | Velocità 0.5-2.5 m/min"),
    ("Saldatura TIG Orbitale", "Diametri 6-323mm | EN ISO 15614 | GMP certificata"),
    ("Pressa Piegatrice 650T", "Forza 650T | Lunghezza 6000mm | Spessore 60mm"),
    ("Sabbiatrice Automatica", "Tunnel 2000×2000mm | Grado Sa 2.5"),
    ("Tornitura CNC Ø2000", "Diametro max 2000mm | Lunghezza 8000mm"),
    ("Gru a Ponte 20T", "Portata 20T | Luce 24m | Altezza 10m"),
]

for equip, specs in equipment_data:
    row_cells = table.add_row().cells
    row_cells[0].text = equip
    row_cells[1].text = specs

doc.add_page_break()

# ===== PORTFOLIO =====
h1 = doc.add_heading("PORTFOLIO PROGETTI", level=1)
h1_format = h1.runs[0]
h1_format.font.color.rgb = NAVY

projects = [
    ("Struttura Portante Turbina Eolica (2024)", "Acciaio S355 | EN 1090 EXC3 | Verniciatura C5-M | Nord Europa"),
    ("Skid Processo Farmaceutico GMP (2024)", "AISI 316L | Piping sanitario | IQ/OQ | Svizzera"),
    ("Piping Industriale Impianto Energia (2023)", "Acciaio carbonio e inox | Pressioni 40 bar | Germania"),
    ("Moduli Prefabbricati Oil & Gas (2023)", "Struttura modulare S275 | PED e ATEX | Benelux"),
    ("Carpenteria Pesante Infrastrutture (2022)", "Travi HEB | Bullonature classe 8.8 | Italia"),
    ("Rack Piping Alta Pressione (2022)", "Tubi P91 | TIG orbitale | Controlli RT+UT 100% | Francia"),
    ("Costruzione Navale - Scafo Acciaio (2023)", "Laminati fino a 100mm | Saldatura sottoacqua | Mediterraneo"),
    ("Tubazioni Coibentate Termoelettrico (2024)", "Acciaio P11 | Coibentazione ceramica 50mm | Pressioni 100 bar"),
]

for proj_name, proj_info in projects:
    p = doc.add_paragraph()
    run = p.add_run(proj_name)
    run.font.bold = True
    p.add_run(f"\n{proj_info}")

doc.add_page_break()

# ===== QUALITÀ E NORMATIVE =====
h1 = doc.add_heading("QUALITÀ E NORMATIVE", level=1)
h1_format = h1.runs[0]
h1_format.font.color.rgb = NAVY

h2 = doc.add_heading("Certificazioni", level=2)
h2_format = h2.runs[0]
h2_format.font.color.rgb = BLUE

doc.add_paragraph("ISO 9001:2015 - Qualità e processi")
doc.add_paragraph("EN 9100:2018 - Aerospaziale (applicabile a settori industriali critici)")

h2 = doc.add_heading("Normative di Riferimento", level=2)
h2_format = h2.runs[0]
h2_format.font.color.rgb = BLUE

norms = [
    "EN 1090 - Carpenteria metallica strutturale",
    "EN ISO 3834 - Qualità saldature",
    "EN ISO 15614 - Qualificazione procedimenti di saldatura",
    "ASME IX - Qualificazione saldatori",
    "PED 2014/68/UE - Attrezzature in pressione",
    "EN ISO 9712 - Personale NDT certificato",
]

for norm in norms:
    doc.add_paragraph(norm)

# ===== CONTATTI =====
h2 = doc.add_heading("CONTATTI", level=2)
h2_format = h2.runs[0]
h2_format.font.color.rgb = BLUE

p = doc.add_paragraph()
run = p.add_run("Email: ")
run.font.bold = True
p.add_run("info@evolutionengineering.eu")

p = doc.add_paragraph()
run = p.add_run("Sito Web: ")
run.font.bold = True
p.add_run("www.evolutionengineering.eu")

# ===== FOOTER =====
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(48)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("© 2025 Evolution Engineering & Services S.r.l. — Tutti i diritti riservati\nDocumento generato automaticamente | 02 Aprile 2025")
run.font.size = Pt(9)
run.font.color.rgb = GRAY

# Salva il documento
doc.save("sito_evolution_engineering.docx")
print("✅ File Word creato: sito_evolution_engineering.docx")
